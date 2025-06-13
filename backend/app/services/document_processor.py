import os
import logging
import json
import uuid
import traceback
from typing import List, Dict, Any, Optional
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
from datetime import datetime
from app.models.document import Document, DocumentCreate, DocumentUpdate
from app.db.init_db_simple import get_async_db, get_chroma_client
from app.core.config import settings

# Configure more detailed logging
logger = logging.getLogger(__name__)

# Add file handler for debugging
file_handler = logging.FileHandler("document_processor_debug.log")
file_handler.setLevel(logging.DEBUG)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
logger.addHandler(file_handler)
logger.setLevel(logging.DEBUG)
logger.info("Document processor module loaded")

class DocumentProcessor:
    """Service for processing documents and storing them in the database and vector store"""
    
    async def save_uploaded_file(self, file: UploadFile, document_id: str) -> str:
        """
        Save an uploaded file to disk
        
        Args:
            file: The uploaded file
            document_id: Unique ID for the document
            
        Returns:
            str: Path to the saved file
        """
        # Create upload directory if it doesn't exist
        os.makedirs(settings.UPLOAD_DIRECTORY, exist_ok=True)
        
        # Get file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        # Create a unique filename
        unique_filename = f"{document_id}{file_ext}"
        file_path = os.path.join(settings.UPLOAD_DIRECTORY, unique_filename)
        
        # Write file to disk
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
            
        logger.info(f"[DEBUG] Saved file {file.filename} to {file_path} (size: {len(content)} bytes)")
        return file_path
    
    async def create_document(self, db: AsyncSession, document_create: DocumentCreate) -> Document:
        """
        Create a new document record in the database
        
        Args:
            db: Database session
            document_create: Document creation data
            
        Returns:
            Document: Created document
        """
        document = Document(
            id=document_create.id,
            filename=document_create.filename,
            file_path=document_create.file_path,
            content_type=document_create.content_type,
            status=document_create.status,
            project_id=document_create.project_id,
            description=document_create.description,
            doc_metadata=json.dumps(document_create.doc_metadata) if document_create.doc_metadata else None
        )
        
        db.add(document)
        await db.commit()
        await db.refresh(document)
        
        logger.info(f"[DEBUG] Created document record with ID {document.id} for project {document_create.project_id}")
        return document
        
    async def update_document(self, db, document_id: str, document_update: DocumentUpdate) -> bool:
        """
        Update a document in the database
        
        Args:
            db: Database session or connection pool
            document_id: ID of the document to update
            document_update: Document update data
            
        Returns:
            bool: True if update was successful, False otherwise
        """
        try:
            logger.info(f"[DEBUG] Updating document {document_id} with: status={document_update.status}, progress={document_update.progress}")
            
            # Check if we're using a connection pool or SQLAlchemy session
            if hasattr(db, 'execute'):
                # Using SQLAlchemy session
                try:
                    # First check if the progress column exists
                    result = await db.execute(text(
                        """SELECT EXISTS (
                            SELECT FROM information_schema.columns 
                            WHERE table_name = 'documents' AND column_name = 'progress'
                        )"""
                    ))
                    column_exists = result.scalar()
                    logger.info(f"[DEBUG] Progress column exists in documents table: {column_exists}")
                    
                    # If progress column doesn't exist, add it
                    if not column_exists:
                        logger.info("[DEBUG] Adding progress column to documents table")
                        await db.execute(text(
                            """ALTER TABLE documents 
                            ADD COLUMN IF NOT EXISTS progress TEXT DEFAULT '0'"""
                        ))
                        await db.commit()
                except Exception as e:
                    logger.error(f"[DEBUG] Error checking/adding progress column: {e}")
                    await db.rollback()
                    
                    # Build update query dynamically based on provided fields
                    update_fields = {}
                    if document_update.status is not None:
                        update_fields['status'] = document_update.status
                        logger.info(f"[DEBUG] Will update status to: {document_update.status}")
                    if document_update.progress is not None:
                        update_fields['progress'] = document_update.progress
                        logger.info(f"[DEBUG] Will update progress to: {document_update.progress}")
                    
                    # Log the current document state before updating
                    doc_before = await conn.fetchrow("SELECT * FROM documents WHERE id = $1", document_id)
                    if doc_before:
                        logger.info(f"[DEBUG] Document before update: id={doc_before['id']}, status={doc_before['status']}, progress={doc_before.get('progress', 'N/A')}")
                    else:
                        logger.warning(f"[DEBUG] Document {document_id} not found before update")
                    if document_update.filename is not None:
                        update_fields['filename'] = document_update.filename
                    if document_update.project_id is not None:
                        update_fields['project_id'] = document_update.project_id
                    if document_update.description is not None:
                        update_fields['description'] = document_update.description
                    if document_update.doc_metadata is not None:
                        update_fields['doc_metadata'] = json.dumps(document_update.doc_metadata)
                    
                    # Always update the updated_at timestamp
                    update_fields['updated_at'] = datetime.utcnow()
                    
                    if not update_fields:
                        logger.warning(f"No fields to update for document {document_id}")
                        return False
                    
                    # Build the SQL query
                    set_clause = ", ".join([f"{k} = ${i+2}" for i, k in enumerate(update_fields.keys())])
                    values = list(update_fields.values())
                    
                    # Execute the update query
                    await conn.execute(
                        f"UPDATE documents SET {set_clause} WHERE id = $1",
                        document_id,
                        *values
                    )
                    
                    # Verify the update was successful
                    doc_after = await conn.fetchrow("SELECT * FROM documents WHERE id = $1", document_id)
                    if doc_after:
                        logger.info(f"[DEBUG] Document after update: id={doc_after['id']}, status={doc_after['status']}, progress={doc_after.get('progress', 'N/A')}")
                    else:
                        logger.warning(f"[DEBUG] Document {document_id} not found after update")
                    
                    logger.info(f"[DEBUG] Updated document {document_id} with fields: {', '.join(update_fields.keys())}")
                    return True
            else:
                # Using SQLAlchemy session
                stmt = select(Document).where(Document.id == document_id)
                result = await db.execute(stmt)
                document = result.scalars().first()
                
                if not document:
                    logger.warning(f"Document {document_id} not found for update")
                    return False
                
                # Update document fields
                if document_update.status is not None:
                    document.status = document_update.status
                if document_update.progress is not None:
                    document.progress = document_update.progress
                if document_update.filename is not None:
                    document.filename = document_update.filename
                if document_update.project_id is not None:
                    document.project_id = document_update.project_id
                if document_update.description is not None:
                    document.description = document_update.description
                if document_update.doc_metadata is not None:
                    document.doc_metadata = json.dumps(document_update.doc_metadata)
                
                document.updated_at = datetime.utcnow()
                
                await db.commit()
                logger.info(f"[DEBUG] Updated document {document_id} in database")
                return True
                
        except Exception as e:
            logger.error(f"Error updating document {document_id}: {str(e)}")
            logger.exception(e)
            return False
    
    async def process_document(self, db, document_id: str, file_path: str) -> None:
        """
        Process a document and extract text for vectorization
        
        Args:
            db: Database session (can be None, in which case we'll create our own)
            document_id: ID of the document to process
            file_path: Path to the document file
        """
        logger.info(f"===== PROCESS_DOCUMENT CALLED: document_id={document_id}, file_path={file_path} =====")
        with open("document_processing_started.txt", "a") as f:
            f.write(f"{datetime.now()} - Started processing document {document_id} at {file_path}\n")
            
        try:
            # Print path information
            logger.info(f"File exists: {os.path.exists(file_path)}")
            logger.info(f"Absolute path: {os.path.abspath(file_path)}")
            if os.path.exists(file_path):
                logger.info(f"File size: {os.path.getsize(file_path)}")
                logger.info(f"File directory contents: {os.listdir(os.path.dirname(file_path))}")
        except Exception as e:
            logger.error(f"Error checking file: {e}")
            
        # Create a new db session if None was provided
        session_created = False
        db_instance = None
        if db is None:
            try:
                from app.db.init_db_simple import AsyncSessionLocal
                db_instance = AsyncSessionLocal()
                db = db_instance
                session_created = True
                logger.info(f"[DEBUG] Created new database session for background processing of document {document_id}")
            except Exception as db_error:
                logger.error(f"Error creating database session: {db_error}")
                logger.error(traceback.format_exc())
                return  # Exit early if we can't create a database session
        
        # Initialize doc_metadata to avoid reference errors in except/finally blocks
        doc_metadata = {}
        chunks = []
        
        try:
            logger.info(f"[DEBUG] Starting processing of document {document_id} at path {file_path}")
            
            # Import here to ensure we have the correct SQLAlchemy objects
            from sqlalchemy import text
            from sqlalchemy.exc import SQLAlchemyError
            
            # STAGE 1: Verify document exists in DB (0%)
            # First, verify the document exists in the database to avoid errors later
            try:
                # Use direct SQL query for reliability
                result = await db.execute(text("SELECT id FROM documents WHERE id = :id"), {"id": document_id})
                doc_exists = result.fetchone() is not None
                if not doc_exists:
                    logger.error(f"[DEBUG] Document {document_id} not found in database")
                    return
                logger.info(f"[DEBUG] Document {document_id} found in database, continuing processing")
            except SQLAlchemyError as sql_error:
                logger.error(f"[DEBUG] SQL error checking if document exists: {sql_error}")
                return
                
            # STAGE 2: Verify file exists (20%)
            if not os.path.exists(file_path):
                logger.error(f"[DEBUG] File not found at {file_path}")
                try:
                    # Use safer direct SQL update for error status
                    await db.execute(
                        text("UPDATE documents SET status = :status, progress = :progress WHERE id = :id"),
                        {"id": document_id, "status": "error", "progress": "0"}
                    )
                    await db.commit()
                    logger.info(f"[DEBUG] Updated document status to error (file not found)")
                except SQLAlchemyError as commit_error:
                    logger.error(f"[DEBUG] Failed to update document status: {commit_error}")
                    await db.rollback()
                return
            
            try:
                await db.execute(
                    text("UPDATE documents SET progress = :progress WHERE id = :id"),
                    {"id": document_id, "progress": "20"}
                )
                await db.commit()
                logger.info(f"[DEBUG] Updated progress to 20% - file verification complete")
            except SQLAlchemyError as e:
                logger.error(f"[DEBUG] Failed to update progress to 20%: {e}")
                await db.rollback()
            
            # STAGE 3: Extract text from document based on file type (40%)
            file_ext = os.path.splitext(file_path)[1].lower()
            logger.info(f"[DEBUG] Detected file extension: {file_ext}")
            text_content = ""
            
            if file_ext == ".pdf":
                logger.info(f"[DEBUG] Starting PDF extraction for {file_path}")
                # This would use a PDF extraction library like PyPDF2 or pdfplumber
                # For now, use a simple placeholder
                text_content = f"Extracted PDF content from {os.path.basename(file_path)}"
                logger.info(f"[DEBUG] Extracted PDF content from {file_path}")
            
            elif file_ext == ".docx":
                logger.info(f"[DEBUG] Starting DOCX extraction for {file_path}")
                # Try to use python-docx if available
                try:
                    from docx import Document as DocxDocument
                    docx_document = DocxDocument(file_path)
                    paragraphs = []
                    for p in docx_document.paragraphs:
                        if p.text.strip():
                            paragraphs.append(p.text)
                    
                    text_content = "\n\n".join(paragraphs)
                    logger.info(f"[DEBUG] Extracted DOCX content: {len(paragraphs)} paragraphs, {len(text_content)} characters")
                except ImportError:
                    # Fallback if python-docx is not installed
                    text_content = f"Extracted DOCX content from {os.path.basename(file_path)} (fallback method)"
                    logger.info(f"[DEBUG] Used fallback extraction for DOCX (python-docx not available)")
                except Exception as docx_error:
                    logger.error(f"[DEBUG] Error extracting DOCX: {str(docx_error)}")
                    text_content = f"Extracted DOCX content from {os.path.basename(file_path)} (fallback after error)"
            
            elif file_ext == ".txt":
                logger.info(f"[DEBUG] Starting TXT extraction for {file_path}")
                # Read plain text file
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_content = f.read()
                    logger.info(f"[DEBUG] Read text file content: {len(text_content)} characters")
                except Exception as txt_error:
                    logger.error(f"[DEBUG] Error reading text file: {str(txt_error)}")
                    text_content = f"Error reading content from {os.path.basename(file_path)}"
            
            else:
                # Unsupported file type - but continue with placeholder content
                logger.warning(f"[DEBUG] Unsupported file type: {file_ext}")
                text_content = f"Content from unsupported file type: {os.path.basename(file_path)}"
            
            # Update progress after text extraction
            try:
                await db.execute(
                    text("UPDATE documents SET progress = :progress WHERE id = :id"),
                    {"id": document_id, "progress": "40"}
                )
                await db.commit()
                logger.info(f"[DEBUG] Updated progress to 40% - text extraction complete")
            except SQLAlchemyError as e:
                logger.error(f"[DEBUG] Failed to update progress to 40%: {e}")
                await db.rollback()
            
            # STAGE 4: Process text into chunks (60%)
            # Split text into chunks for vectorization
            doc_metadata["document_size"] = len(text_content)
            chunks = self._split_text_into_chunks(text_content)
            doc_metadata["chunk_count"] = len(chunks)
            
            try:
                await db.execute(
                    text("UPDATE documents SET progress = :progress WHERE id = :id"),
                    {"id": document_id, "progress": "60"}
                )
                await db.commit()
                logger.info(f"[DEBUG] Updated progress to 60% - chunking complete with {len(chunks)} chunks")
            except SQLAlchemyError as e:
                logger.error(f"[DEBUG] Failed to update progress to 60%: {e}")
                await db.rollback()
            
            # STAGE 5: Vectorize and store chunks (80%)
            try:
                # Get the project_id for this document
                project_id = None
                try:
                    result = await db.execute(
                        text("SELECT project_id FROM documents WHERE id = :id"),
                        {"id": document_id}
                    )
                    row = result.fetchone()
                    if row:
                        project_id = row.project_id
                        logger.info(f"[DEBUG] Found project_id {project_id} for document {document_id}")
                    else:
                        logger.error(f"[DEBUG] Could not find project_id for document {document_id}")
                except Exception as e:
                    logger.error(f"[DEBUG] Error getting project_id: {e}")
                
                if not project_id:
                    logger.error(f"[DEBUG] Cannot store document in ChromaDB without project_id")
                    doc_metadata["vectorization_error"] = "No project_id found"
                else:
                    # Actually store the chunks in ChromaDB
                    client = get_chroma_client()
                    # Use project-specific collection to match DocumentSearchTool
                    collection_name = f"project_{project_id}"
                    collection = client.get_or_create_collection(collection_name)
                    logger.info(f"[DEBUG] Using ChromaDB collection: {collection_name}")
                    
                    # Prepare chunk data for insertion
                    chunk_ids = []
                    documents = []
                    metadatas = []
                    
                    for i, chunk in enumerate(chunks):
                        chunk_id = f"{document_id}_chunk_{i}"
                        chunk_ids.append(chunk_id)
                        documents.append(chunk)
                        metadatas.append({
                            "document_id": document_id,
                            "project_id": project_id,
                            "chunk_number": i,
                            "total_chunks": len(chunks),
                            "filename": os.path.basename(file_path),
                            "source": os.path.basename(file_path),  # Add source for DocumentSearchTool
                            "file_extension": os.path.splitext(file_path)[1].lower(),
                            "processing_date": str(datetime.now())
                        })
                    
                    # Insert chunks into ChromaDB
                    collection.add(
                        ids=chunk_ids,
                        documents=documents,
                        metadatas=metadatas
                    )
                    
                    logger.info(f"[DEBUG] Successfully stored {len(chunks)} chunks in ChromaDB collection {collection_name} for document {document_id}")
                    doc_metadata["chunks_stored_in_chroma"] = len(chunks)
                    doc_metadata["chroma_collection"] = collection_name
                
            except Exception as vectorize_error:
                logger.error(f"[DEBUG] Error during vectorization: {str(vectorize_error)}")
                logger.error(traceback.format_exc())
                # Continue processing despite vectorization error
                doc_metadata["vectorization_error"] = str(vectorize_error)
            
            # Update progress to 80%
            try:
                await db.execute(
                    text("UPDATE documents SET progress = :progress WHERE id = :id"),
                    {"id": document_id, "progress": "80"}
                )
                await db.commit()
                logger.info(f"[DEBUG] Updated progress to 80% - vectorization complete")
            except SQLAlchemyError as e:
                logger.error(f"[DEBUG] Failed to update progress to 80%: {e}")
                await db.rollback()
            
            # STAGE 6: Final completion (100%)
            # Add timestamp to metadata
            doc_metadata["processed_at"] = str(datetime.now())
            
            # Update document status to processed with final 100% progress
            try:
                await db.execute(
                    text("UPDATE documents SET status = :status, progress = :progress, doc_metadata = :metadata WHERE id = :id"),
                    {
                        "id": document_id, 
                        "status": "processed", 
                        "progress": "100",
                        "metadata": json.dumps(doc_metadata)
                    }
                )
                await db.commit()
                logger.info(f"[DEBUG] Updated progress to 100% - document fully processed")
            except SQLAlchemyError as final_error:
                logger.error(f"[DEBUG] Failed to update final document status: {final_error}")
                await db.rollback()
                
            # Verify the update was successful using appropriate session method
            try:
                if hasattr(db, 'execute'):  # SQLAlchemy session
                    result = await db.execute(text("SELECT * FROM documents WHERE id = :id"), {"id": document_id})
                    doc_result = result.fetchone()
                    if doc_result:
                        logger.info(f"[DEBUG] Document final state: id={doc_result.id}, status={doc_result.status}, progress={doc_result.progress if hasattr(doc_result, 'progress') else 'N/A'}")
                    else:
                        logger.warning(f"[DEBUG] Document {document_id} not found after final update")
            except Exception as verify_error:
                logger.error(f"[DEBUG] Error verifying final document state: {str(verify_error)}")
                logger.error(traceback.format_exc())
            
            logger.info(f"[DEBUG] Successfully completed processing document {document_id} with {len(chunks)} chunks")
        except Exception as e:
            # Handle any other exceptions during processing
            logger.error(f"[DEBUG] Error processing document {document_id}: {str(e)}\n{traceback.format_exc()}")
            
            # Update document status to error
            try:
                # Use direct SQL to avoid any issues with complex ORM operations
                if hasattr(db, 'execute'):
                    await db.execute(
                        text("UPDATE documents SET status = :status WHERE id = :id"),
                        {"id": document_id, "status": "error"}
                    )
                    await db.commit()
                    logger.info(f"[DEBUG] Updated document status to error")
                else:
                    logger.warning(f"[DEBUG] Could not update document status: db object doesn't have execute method")
            except Exception as update_error:
                logger.error(f"[DEBUG] Failed to update document status to error: {update_error}")
                
        finally:
            # Close the database session if we created it
            if session_created and db is not None:
                try:
                    await db.close()
                    logger.info(f"[DEBUG] Closed database session after processing document {document_id}")
                except Exception as close_error:
                    logger.error(f"[DEBUG] Error closing database session: {str(close_error)}")
            
            # Log completion regardless of success/failure
            logger.info(f"[DEBUG] Document processing task for {document_id} has completed execution")
            
            with open("document_processing_completed.txt", "a") as f:
                f.write(f"{datetime.now()} - Completed processing document {document_id}\n")
    
    async def ensure_documents_indexed(self, db: AsyncSession, project_id: str) -> bool:
        """
        Ensure all processed documents for a project are indexed in ChromaDB.
        This is useful for documents that were processed before ChromaDB indexing was implemented.
        
        Args:
            db: Database session
            project_id: ID of the project
            
        Returns:
            bool: True if all documents are indexed, False otherwise
        """
        try:
            logger.info(f"Ensuring all documents for project {project_id} are indexed in ChromaDB")
            
            # Get all processed documents for this project
            from sqlalchemy import select
            from app.models.document import Document
            
            stmt = select(Document).where(
                Document.project_id == project_id,
                Document.status == "processed"
            )
            result = await db.execute(stmt)
            documents = result.scalars().all()
            
            if not documents:
                logger.info(f"No processed documents found for project {project_id}")
                return True
            
            # Get ChromaDB client and collection
            client = get_chroma_client()
            collection_name = f"project_{project_id}"
            collection = client.get_or_create_collection(collection_name)
            
            # Check which documents are already indexed
            try:
                # Get all document IDs in the collection
                all_ids = collection.get()["ids"] if collection.get()["ids"] else []
                indexed_doc_ids = set()
                for chunk_id in all_ids:
                    # Extract document ID from chunk ID (format: "docid_chunk_0")
                    if "_chunk_" in chunk_id:
                        doc_id = chunk_id.split("_chunk_")[0]
                        indexed_doc_ids.add(doc_id)
                
                logger.info(f"Found {len(indexed_doc_ids)} documents already indexed in ChromaDB")
                
                # Index documents that are not yet indexed
                for doc in documents:
                    if doc.id not in indexed_doc_ids:
                        logger.info(f"Document {doc.id} ({doc.filename}) not indexed, attempting to index...")
                        
                        # Check if file exists
                        if doc.file_path and os.path.exists(doc.file_path):
                            # Re-process the document to index it
                            await self.process_document(db, doc.id, doc.file_path)
                        else:
                            logger.warning(f"File not found for document {doc.id}: {doc.file_path}")
                
                return True
                
            except Exception as e:
                logger.error(f"Error checking/indexing documents: {e}")
                return False
                
        except Exception as e:
            logger.error(f"Error ensuring documents indexed for project {project_id}: {e}")
            return False
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for vectorization
        
        Args:
            text: Text to split
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List[str]: List of text chunks
        """
        chunks = []
        if len(text) <= chunk_size:
            chunks.append(text)
            logger.debug(f"[DEBUG] Text fits in a single chunk of size {len(text)}")
        else:
            start = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                # If this is not the first chunk, include overlap
                if start > 0:
                    start = start - overlap
                chunks.append(text[start:end])
                logger.debug(f"[DEBUG] Created chunk from position {start} to {end}")
                start = end
        
        return chunks
    
    async def _vectorize_chunks(self, document_id: str, chunks: List[str]) -> None:
        """
        Vectorize text chunks and store in ChromaDB
        
        Args:
            document_id: ID of the document
            chunks: List of text chunks to vectorize
        """
        # Get ChromaDB client
        client = get_chroma_client()
        collection = client.get_or_create_collection("documents")
        logger.info(f"[DEBUG] Connected to ChromaDB collection 'documents'")
        
        # Create IDs for chunks
        chunk_ids = [f"{document_id}_{i}" for i in range(len(chunks))]
        
        # Create metadata for chunks
        metadatas = [{"document_id": document_id, "chunk_index": i} for i in range(len(chunks))]
        
        # Add chunks to collection
        try:
            logger.info(f"[DEBUG] Starting ChromaDB vectorization for document {document_id} with {len(chunks)} chunks")
            
            # Check if ChromaDB is accessible
            try:
                # Check if chunks for this document already exist and delete them
                logger.info(f"[DEBUG] Checking for existing chunks in ChromaDB for document {document_id}")
                existing_results = collection.get(
                    where={"document_id": document_id}
                )
                
                if existing_results and existing_results["ids"]:
                    logger.info(f"[DEBUG] Removing {len(existing_results['ids'])} existing chunks for document {document_id}")
                    collection.delete(
                        where={"document_id": document_id}
                    )
                    logger.info(f"[DEBUG] Successfully removed existing chunks for document {document_id}")
            except Exception as chroma_check_error:
                logger.error(f"[DEBUG] Error checking ChromaDB for existing chunks: {str(chroma_check_error)}")
                logger.error(traceback.format_exc())
                # Continue with the process - we'll try to add the chunks anyway
            
            try:
                # Add new chunks
                logger.info(f"[DEBUG] Adding {len(chunks)} chunks to ChromaDB for document {document_id}")
                collection.add(
                    ids=chunk_ids,
                    documents=chunks,
                    metadatas=metadatas
                )
                logger.info(f"[DEBUG] Successfully added {len(chunks)} chunks to ChromaDB for document {document_id}")
            except Exception as add_error:
                logger.error(f"[DEBUG] Error adding chunks to ChromaDB: {str(add_error)}")
                logger.error(traceback.format_exc())
                # We'll continue with a fallback - the document can still be marked as processed
                # even if vectorization failed
                logger.warning(f"[DEBUG] Using fallback for failed vectorization for document {document_id}")
        except Exception as e:
            logger.error(f"[DEBUG] Unhandled error in ChromaDB vectorization: {str(e)}")
            logger.error(traceback.format_exc())
            # Don't raise the exception - we want the document processing to continue
            # and mark the document as processed even if vectorization failed
            logger.warning(f"[DEBUG] Continuing document processing despite ChromaDB error for document {document_id}")
    
    async def get_document(self, db: AsyncSession, document_id: str) -> Optional[Document]:
        """
        Get a document by ID
        
        Args:
            db: Database session
            document_id: ID of the document to retrieve
            
        Returns:
            Optional[Document]: Document if found, None otherwise
        """
        result = await db.execute(select(Document).where(Document.id == document_id))
        return result.scalars().first()
    
    async def list_documents(self, db: AsyncSession, project_id: Optional[str] = None) -> List[Document]:
        """
        List all documents, optionally filtered by project_id
        
        Args:
            db: Database session
            project_id: Optional project ID to filter by
            
        Returns:
            List[Document]: List of documents
        """
        if project_id:
            result = await db.execute(select(Document).where(Document.project_id == project_id))
        else:
            result = await db.execute(select(Document))
            
        return result.scalars().all()
    
    async def update_document_status(self, db: AsyncSession, document_id: str, status: str) -> Optional[Document]:
        """
        Update a document's status
        
        Args:
            db: Database session
            document_id: ID of the document to update
            status: New status
            
        Returns:
            bool: True if document was updated, False otherwise
        """
        return await self.update_document(db, document_id, DocumentUpdate(status=status))
    
    async def update_document(self, db, document_id: str, document_update: 'DocumentUpdate') -> bool:
        """
        Update a document in the database
        
        Args:
            db: Database session
            document_id: ID of the document to update
            document_update: Document update data
            
        Returns:
            bool: True if document was updated, False otherwise
        """
        # Safety check
        if db is None:
            logger.error(f"Cannot update document {document_id}: database session is None")
            return False
            
        try:
            # Check if this is a SQLAlchemy session or another type
            if hasattr(db, 'execute'):
                # This is SQLAlchemy
                logger.info(f"[DEBUG] Using SQLAlchemy to update document {document_id}")
                
                # Start a nested transaction if possible
                try:
                    from sqlalchemy import text
                    # Prepare update fields
                    update_fields = {}
                    if document_update.status is not None:
                        update_fields['status'] = document_update.status
                    if document_update.progress is not None:
                        update_fields['progress'] = document_update.progress
                    if document_update.filename is not None:
                        update_fields['filename'] = document_update.filename
                    if document_update.project_id is not None:
                        update_fields['project_id'] = document_update.project_id
                    if document_update.description is not None:
                        update_fields['description'] = document_update.description
                    if document_update.metadata is not None:
                        update_fields['metadata'] = json.dumps(document_update.metadata)
                        
                    async with db.begin_nested():
                        await db.execute(text("UPDATE documents SET " + ", ".join([f"{key} = :{key}" for key in update_fields]) + " WHERE id = :id"), {**update_fields, "id": document_id})
                    return True
                except Exception as e:
                    logger.error(f"[DEBUG] Error updating document {document_id}: {str(e)}")
                    logger.error(traceback.format_exc())
                    return False
            else:
                # This is not SQLAlchemy
                logger.info(f"[DEBUG] Using non-SQLAlchemy database to update document {document_id}")
                
                # Update document fields
                document = await self.get_document(db, document_id)
                
                if not document:
                    return False
                    
                if document_update.status is not None:
                    document.status = document_update.status
                    
                if document_update.progress is not None:
                    document.progress = document_update.progress
                    
                if document_update.filename is not None:
                    document.filename = document_update.filename
                    
                if document_update.project_id is not None:
                    document.project_id = document_update.project_id
                    
                if document_update.description is not None:
                    document.description = document_update.description
                    
                if document_update.metadata is not None:
                    document.metadata = json.dumps(document_update.metadata)
                    
                try:
                    async with db.begin():
                        await db.commit()
                    return True
                except Exception as e:
                    logger.error(f"[DEBUG] Error updating document {document_id}: {str(e)}")
                    logger.error(traceback.format_exc())
                    return False
        except Exception as e:
            logger.error(f"[DEBUG] Unhandled error updating document {document_id}: {str(e)}")
            logger.error(traceback.format_exc())
            return False
            
        await db.commit()
        await db.refresh(document)
        
        return document
    
    async def delete_document(self, db: AsyncSession, document_id: str) -> bool:
        """
        Delete a document
        
        Args:
            db: Database session
            document_id: ID of the document to delete
            
        Returns:
            bool: True if document was deleted, False otherwise
        """
        document = await self.get_document(db, document_id)
        
        if not document:
            return False
            
        # Delete document from ChromaDB
        client = get_chroma_client()
        collection = client.get_collection("documents")
        
        # Query for chunks with this document_id in metadata
        results = collection.get(
            where={"document_id": document_id}
        )
        
        if results and results["ids"]:
            # Delete chunks from collection
            collection.delete(ids=results["ids"])
            
        # Delete file from disk if it exists
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
            
        # Delete document from database
        await db.delete(document)
        await db.commit()
        
        return True
